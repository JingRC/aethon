"""
Word 文档生成器 — 小红书发布套装

生成一份精美的 .docx 文件，包含：
  1. 爆款标题
  2. 正文描述（1000字以内）
  3. AI 快讯链接 + 摘要速览表
"""

import logging
from datetime import datetime

logger = logging.getLogger(__name__)


def generate_xhs_docx(ai_news: list[dict],
                      xhs_content: dict,
                      output_path: str) -> str:
    """Generate a Word document with XHS publishing kit."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE

    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # -- Styles --
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # ── Title ──
    title = xhs_content.get("title", "今日AI快讯")
    h = doc.add_heading(title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = RGBColor(0xFF, 0x33, 0x66)

    # ── Date ──
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(datetime.now().strftime("%Y年%m月%d日"))
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()  # spacer

    # ── Body ──
    body = xhs_content.get("body", "")
    doc.add_heading("📝 正文描述", level=2)
    body_para = doc.add_paragraph(body[:1000])  # 1000字限制
    body_para.paragraph_format.line_spacing = 1.6

    # ── Hashtags ──
    hashtags = xhs_content.get("hashtags", [])
    if hashtags:
        doc.add_heading("🏷️ 话题标签", level=2)
        tags_para = doc.add_paragraph("  ".join(hashtags))
        for run in tags_para.runs:
            run.font.color.rgb = RGBColor(0x00, 0x88, 0xFF)

    # ── News Links & Summaries ──
    doc.add_heading("🔗 AI快讯链接 & 摘要速览", level=2)
    for i, item in enumerate(ai_news[:10], 1):
        title_text = item.get("title", "")
        source = item.get("source", "")
        summary = item.get("summary_zh", "")
        url = item.get("url", "")
        score = int(item.get("score", 3))
        stars = "★" * score + "☆" * (5 - score)

        # Item heading
        h3 = doc.add_heading(f"#{i} {stars}  {title_text}", level=3)
        for run in h3.runs:
            run.font.size = Pt(12)

        # Source
        src_para = doc.add_paragraph()
        run = src_para.add_run(f"来源: {source}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        # Summary
        if summary:
            doc.add_paragraph(summary, style='List Bullet')

        # URL
        if url:
            url_para = doc.add_paragraph()
            run = url_para.add_run(url)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x66, 0xCC)

        doc.add_paragraph()  # spacer

    # ── Footer ──
    doc.add_paragraph("─" * 40)
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run("🤖 由每日AI日报自动生成 | GitHub Actions 云端运行")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.save(output_path)
    logger.info(f"[DOCX] Saved: {output_path}")
    return output_path
