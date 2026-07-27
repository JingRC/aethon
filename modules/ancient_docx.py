"""
Word 文档生成器 — 历史故事小红书发布套装

生成一份 .docx 文件，包含：
  1. 爆款标题
  2. 正文描述
  3. 话题标签
  4. 10 则故事速览表
"""

import logging
from datetime import datetime

logger = logging.getLogger(__name__)


def build_history_xhs_prompt(stories: list[dict], count: int = 10) -> str:
    """构建历史故事 XHS 爆款文案 LLM prompt（v2：注入 7 大爆款标题公式）。"""
    story_blocks = []
    for i, s in enumerate(stories[:count], 1):
        blocks = (
            f"{i}. 标题：{s.get('title','')}\n"
            f"   朝代：{s.get('dynasty','')}\n"
            f"   分类：{s.get('category','')}\n"
            f"   故事摘要：{s.get('story_zh','')[:120]}\n"
            f"   启示：{s.get('lesson','')}\n"
        )
        story_blocks.append(blocks)

    return f"""你是一位小红书国学/历史类万粉博主，擅长用年轻化、有网感的语言包装中国传统文化。

请根据以下 {count} 则中国古代历史故事，撰写一篇小红书图文笔记的文案。

## 故事素材
{chr(10).join(story_blocks)}

## 爆款标题公式（必须从以下 7 种中选择一种）

| 公式 | 结构 | 历史赛道示例 |
|------|------|-------------|
| ① 身份定位法 | 所有[A]！把[C]焊脑子里！ | 所有历史迷！把这10个典故焊脑子里！ |
| ② 结果震撼法 | [A]如何用[C]实现[B]？ | 普通人如何用10个成语看透人性？ |
| ③ 痛点反问法 | [B]的人，是不是都被……骗了？ | 喜欢历史的人，是不是都被这些"常识"骗了？ |
| ④ 秘密揭露法 | [A]不会告诉你的[C] | 历史课本不会告诉你的10个真相 |
| ⑤ 对比冲突法 | [时间/状态]前 vs 后，[情绪词]！ | 读这些故事前 vs 读完后，谁懂啊这格局！ |
| ⑥ 数字盘点法 | [A][B]？看这[数字]个[C]就够了 | 想变通透？看这10个故事就够了 |
| ⑦ 价值承诺法 | 答应我，[A]试试[C]，[B]开挂！ | 答应我，每天看一个故事，认知直接开挂！ |

**标题硬规则**：
- ≤20 字，6-12 字最优（点击率最高）
- 必须包含数字（如"10个""3则"）或情绪词（谁懂啊/绝了/救命/真香）
- 关键词前置（前4字必须让人知道是历史/成语/典故类内容）
- **禁止**平淡表述："每日典故""历史卡片""XX故事精选"

## 写作要求

### 封面大字 headline（≤8字）
- 用于卡片封面，字极少但要震撼
- 示例："古人卷疯了""看完我跪了""颠覆三观""改变认知"
- 禁止："历史故事""成语典故""每日精选"

### 正文（300-500字）
- **前3句必须抓人**：用一个最颠覆/最震撼/最好笑的故事开场，直接给结论
- 小红书口吻：口语化但不低幼，像在跟朋友分享"你知道吗"级别的冷知识
- 每段 1-2 句，用 emoji 做视觉锚点（🔥📖⚔️🏯💡🤯 等）
- 选 3-4 个最精彩的故事展开，**不要逐条罗列**
- 正文中间埋 1-2 句"金句"（值得截图保存的那种）
- 文末引导互动（"你最想穿越到哪个朝代？""哪个成语你一直用错了？"）

### 话题标签（5-8个）
- 必须全中文，分层配置：
  - 一级（大类）：#历史故事 #国学智慧 #成语典故 #古人智慧
  - 二级（细分）：#每天学点历史 #传统文化 #冷知识 #读书笔记
  - 三级（热点）：#自我提升 #认知升级 #格局打开

## 返回 JSON（只返回 JSON，不要代码块）
{{
  "title": "爆款标题 6-20字",
  "body": "正文 300-500字",
  "hashtags": ["标签1", "标签2", ...],
  "headline": "封面大字 ≤8字"
}}"""


def generate_ancient_docx(stories: list[dict],
                          xhs_copy: dict,
                          output_path: str) -> str:
    """生成历史故事小红书发布文案 Word 文档。"""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # ── 标题 ──
    title_text = xhs_copy.get("title", "历史故事精选")
    h = doc.add_heading(title_text, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)

    # ── 日期 ──
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(datetime.now().strftime("%Y年%m月%d日"))
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()

    # ── 正文 ──
    body = xhs_copy.get("body", "")
    doc.add_heading("📝 正文描述", level=2)
    body_para = doc.add_paragraph(body[:1500])
    body_para.paragraph_format.line_spacing = 1.6

    # ── 话题标签 ──
    hashtags = xhs_copy.get("hashtags", [])
    if hashtags:
        doc.add_heading("🏷️ 话题标签", level=2)
        normalized = [f"#{t.lstrip('#')}" for t in hashtags]
        tags_para = doc.add_paragraph("  ".join(normalized))
        for run in tags_para.runs:
            run.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)

    # ── 封面标题 ──
    headline = xhs_copy.get("headline", "")
    if headline:
        doc.add_heading("🖼️ 封面大字", level=2)
        hl = doc.add_paragraph(headline)
        for run in hl.runs:
            run.font.size = Pt(16)
            run.font.bold = True

    # ── 故事速览表 ──
    doc.add_heading("📜 10则故事速览", level=2)
    for i, s in enumerate(stories[:10], 1):
        title = s.get("title", "")
        dynasty = s.get("dynasty", "")
        category = s.get("category", "")
        lesson = s.get("lesson", "")

        h3 = doc.add_heading(f"#{i}  {dynasty} · {title}  [{category}]", level=3)
        for run in h3.runs:
            run.font.size = Pt(12)

        if lesson:
            p = doc.add_paragraph(f"💡 {lesson}")
            p.paragraph_format.line_spacing = 1.4

        doc.add_paragraph()

    doc.save(output_path)
    logger.info(f"[DOCX] 历史故事文案: {output_path}")
    return output_path


def generate_ancient_docx_fallback(stories: list[dict], output_path: str) -> str:
    """LLM 文案缺失时的 fallback：从故事数据自动生成标题和正文。"""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from datetime import datetime

    count = min(len(stories), 10)
    selected = stories[:count]

    # ── 自动生成标题 ──
    categories = list(set(s.get("category", "") for s in selected if s.get("category")))
    category_str = categories[0] if categories else "历史典故"

    title_templates = [
        f"{count}个改变认知的{category_str}故事",
        f"读完这{count}则{category_str}，我悟了",
        f"历史课本不会告诉你的{count}个{category_str}",
        f"{count}个被误读千年的{category_str}",
        f"普通人不知道的{count}个{category_str}",
        f"答应我，看完这{count}个{category_str}再划走",
        f"🔥 {count}则{category_str}，老祖宗的顶级智慧",
    ]
    auto_title = title_templates[(count + datetime.now().day) % len(title_templates)]

    # ── 自动生成正文 ──
    # 取第 1 个故事作为开场钩子
    intro_story = selected[0] if selected else {}
    LQ = "“"  # 左弯引号 "
    RQ = "”"  # 右弯引号 "
    intro_line = (
        f"你知道吗？{LQ}{intro_story.get('title', '')}{RQ}这个典故背后，藏着一段让人拍案叫绝的历史。"
        if intro_story.get("title") else
        "今天分享10个精彩的历史故事，每一条都让你拍案叫绝。"
    )

    # 选 3-4 个最精彩的故事做摘要
    highlights = []
    for i, s in enumerate(selected[:4]):
        title = s.get("title", "")
        lesson = s.get("lesson", "")
        if lesson:
            highlights.append(f"📖 {title}——{lesson}")
        elif title:
            dynasty = s.get("dynasty", "")
            highlights.append(f"📖 {dynasty}·{title}")

    # 结尾互动
    closings = [
        "哪一个故事最让你意外？评论区告诉我 👇",
        "你最想穿越到哪个朝代？来聊聊～",
        "哪个典故你一直理解错了？评论区见！",
        "收藏起来，每天一个故事提升格局 💪",
    ]
    closing = closings[count % len(closings)]

    auto_body = (
        f"{intro_line}\n\n"
        f"{chr(10).join(highlights)}\n\n"
        f"完整{count}则故事已整理在上方卡片中，向左滑动即可逐张查看 🔥\n\n"
        f"{closing}"
    )

    # ── 自动生成标签 ──
    auto_tags_raw = [
        "历史故事", "国学智慧", category_str, "古人智慧",
        "每天学点历史", "传统文化", "冷知识", "认知升级",
    ]
    auto_tags = [f"#{t.lstrip('#')}" for t in auto_tags_raw]

    # ── 构建文档 ──
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # ── 标题 ──
    h = doc.add_heading(auto_title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)

    # ── 日期 ──
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(datetime.now().strftime("%Y年%m月%d日"))
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()

    # ── 正文 ──
    doc.add_heading("📝 正文描述", level=2)
    body_para = doc.add_paragraph(auto_body)
    body_para.paragraph_format.line_spacing = 1.6

    # ── 话题标签 ──
    doc.add_heading("🏷️ 话题标签", level=2)
    tags_para = doc.add_paragraph("  ".join(auto_tags))
    for run in tags_para.runs:
        run.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)

    # ── 故事速览 ──
    doc.add_heading("📜 故事速览", level=2)
    for i, s in enumerate(selected, 1):
        title = s.get("title", "")
        dynasty = s.get("dynasty", "")
        category = s.get("category", "")
        story_zh = s.get("story_zh", "")
        lesson = s.get("lesson", "")

        h3 = doc.add_heading(f"#{i}  {dynasty} · {title}  [{category}]", level=3)
        for run in h3.runs:
            run.font.size = Pt(12)

        if story_zh:
            p = doc.add_paragraph(story_zh[:300])
            p.paragraph_format.line_spacing = 1.4

        if lesson:
            p = doc.add_paragraph(f"💡 {lesson}")
            p.paragraph_format.line_spacing = 1.4

        doc.add_paragraph()

    doc.save(output_path)
    logger.info(f"[DOCX fallback] 历史故事自动文案: {output_path}")
    return output_path
